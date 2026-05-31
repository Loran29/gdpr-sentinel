"""Text extractor for PDF and Word (.docx) documents.

PDF flow:
  1. pdfplumber  — native text layer (fast, exact)
  2. byte-sweep  — regex over raw PDF stream (synthetic/placeholder PDFs)
  3. OCR         — pytesseract + pdf2image if both are installed

Word flow:
  1. python-docx — paragraph + table cell extraction
"""

from __future__ import annotations

import io
import logging
import re
from dataclasses import dataclass

logger = logging.getLogger(__name__)


@dataclass
class PageText:
    page_number: int
    text: str


def extract_text(file_bytes: bytes, filename: str = "") -> list[PageText]:
    """Dispatch to the correct extractor based on file extension or magic bytes."""
    name_lower = filename.lower()
    if name_lower.endswith(".docx"):
        return _extract_docx(file_bytes)
    # Default: treat as PDF (original behaviour)
    return _extract_pdf(file_bytes)


def _extract_docx(docx_bytes: bytes) -> list[PageText]:
    """Extract text from a Word .docx file using python-docx."""
    try:
        from docx import Document  # type: ignore
        doc = Document(io.BytesIO(docx_bytes))
        lines: list[str] = []

        # Paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                lines.append(text)

        # Tables (e.g. form fields in IT access requests)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    lines.append(row_text)

        full_text = "\n".join(lines)
        if full_text.strip():
            return [PageText(page_number=1, text=full_text)]
    except Exception as exc:  # noqa: BLE001
        logger.warning("python-docx extraction failed: %s", exc)

    return [PageText(page_number=1, text="")]


def _extract_pdf(pdf_bytes: bytes) -> list[PageText]:
    """Extract per-page text from a PDF, with OCR fallback for scanned docs."""
    pages: list[PageText] = []
    try:
        import pdfplumber

        with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
            for i, page in enumerate(pdf.pages, start=1):
                txt = page.extract_text() or ""
                pages.append(PageText(page_number=i, text=txt))
    except Exception as exc:  # noqa: BLE001
        logger.warning("pdfplumber failed; falling back to byte-sweep: %s", exc)

    if any(p.text.strip() for p in pages):
        return pages

    # Fallback 1: byte-sweep for synthetic placeholder PDFs.
    recovered = "\n".join(
        re.findall(r"\(([^()]+)\)\s*Tj", pdf_bytes.decode("latin-1", errors="ignore"))
    )
    if recovered.strip():
        return [PageText(page_number=1, text=recovered)]

    # Fallback 2: OCR via pytesseract + pdf2image (optional dependency).
    ocr_pages = _try_ocr(pdf_bytes)
    if ocr_pages:
        return ocr_pages

    return pages or [PageText(page_number=1, text="")]


def _try_ocr(pdf_bytes: bytes) -> list[PageText]:
    """Attempt OCR on a scanned PDF. Returns empty list if not available."""
    try:
        import pytesseract
        from pdf2image import convert_from_bytes
    except ImportError:
        return []

    try:
        pytesseract.get_tesseract_version()
    except Exception:  # noqa: BLE001
        return []

    try:
        images = convert_from_bytes(pdf_bytes, dpi=200)
        result: list[PageText] = []
        for i, img in enumerate(images, start=1):
            text_en = pytesseract.image_to_string(img, lang="eng") or ""
            try:
                text_de = pytesseract.image_to_string(img, lang="deu") or ""
            except Exception:  # noqa: BLE001
                text_de = ""
            text = text_en if len(text_en) >= len(text_de) else text_de
            result.append(PageText(page_number=i, text=text.strip()))
        if any(p.text for p in result):
            logger.info("OCR extracted text from %d page(s)", len(result))
            return result
    except Exception as exc:  # noqa: BLE001
        logger.warning("OCR failed: %s", exc)
    return []

